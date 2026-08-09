from decimal import Decimal
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.core.exceptions import ValidationError
from django.utils import timezone
from requirements_mgmt.models import Requirement
from .finance_utils import get_fund_balance
from .models import ProcurementCase, NotingSheet, NotingSheetItem, EAS, ConveningOrder
from .forms import ( CaseStageDataForm, ReturnCaseForm, NotingSheetForm, NotingSheetItemFormSet, NotingCFADecisionForm, EASForm, EASCFADecisionForm, EASDocumentUploadForm, ConveningOrderForm, )
from io import BytesIO

from django.http import HttpResponse

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

def compute_eas_autofill(noting_sheet):
    """
    The 8 EAS fields that auto-fetch from their linked NotingSheet,
    read-only on the form:
      file_no          <- noting_sheet.file_no
      eas_id           <- "EAS-" + Requirement ID, e.g. EAS-REQ-2026-00001
      dsc_goods        <- item description (first item's, via
                           NotingSheet.item_name)
      purpose_broad    <- noting_sheet.paragraph_1 (Purport / Subject)
      qty_sanctioned   <- SUM of quantities across all item rows
      amount_sanction  <- noting_sheet.total_amount (all items)
      cost_per_unit    <- amount_sanction / qty_sanctioned (weighted
                           average — NotingSheet supports multiple
                           differently-priced items, so there's no
                           single "the" cost per unit; this keeps
                           qty * cost_per_unit == amount_sanction true)
      sub_details_heads <- the Requirement's Sub Head, stringified

    Called on every GET/POST for both eas_create and eas_edit so the
    displayed/saved values always reflect the current NotingSheet state,
    not a stale snapshot from whenever the EAS was first created.
    """
    requirement = noting_sheet.requirement
    items = noting_sheet.items.all()

    total_qty = sum((item.quantity for item in items), 0)
    total_amount = noting_sheet.total_amount or Decimal("0")
    cost_per_unit = (total_amount / total_qty) if total_qty else Decimal("0")

    sub_head = getattr(requirement, "sub_head", None)

    return {
        "file_no": noting_sheet.file_no,
        "eas_id": f"EAS-{requirement.requirement_id}",
        "dsc_goods": noting_sheet.item_name,
        "purpose_broad": noting_sheet.paragraph_1,
        "qty_sanctioned": total_qty,
        "amount_sanction": total_amount,
        "cost_per_unit": cost_per_unit,
        "sub_details_heads": str(sub_head) if sub_head else "",
    }


# ============================================================
# EXISTING VIEWS — unchanged
# ============================================================

@login_required
def case_detail(request, case_number):
    case = get_object_or_404(ProcurementCase, case_number=case_number)
    history = case.history.select_related('performed_by').order_by('created_at')

    stage_form = CaseStageDataForm(instance=case)
    return_form = ReturnCaseForm()

    if request.method == 'POST':
        action = request.POST.get('action')

        if action == 'save_stage':
            stage_form = CaseStageDataForm(request.POST, instance=case)
            if stage_form.is_valid():
                obj = stage_form.save(commit=False)
                obj.modified_by = request.user
                obj.save()
                messages.success(request, 'Stage data saved.')
                return redirect('case_detail', case_number=case.case_number)

        elif action == 'advance':
            try:
                case.advance(user=request.user)
                messages.success(request, f'Case advanced to {case.get_current_stage_display()}.')
            except ValidationError as e:
                messages.error(request, e.message if hasattr(e, 'message') else str(e))
            return redirect('case_detail', case_number=case.case_number)

        elif action == 'return_case':
            return_form = ReturnCaseForm(request.POST)
            if return_form.is_valid():
                try:
                    case.return_to_stage(
                        user=request.user,
                        target_stage=return_form.cleaned_data['target_stage'],
                        reason=return_form.cleaned_data['reason'],
                    )
                    messages.success(request, f'Case returned to {case.get_current_stage_display()}.')
                except ValidationError as e:
                    messages.error(request, e.message if hasattr(e, 'message') else str(e))
                return redirect('case_detail', case_number=case.case_number)

    context = {
        'case': case,
        'history': history,
        'stage_form': stage_form,
        'return_form': return_form,
        'role': request.user.role,
    }
    return render(request, 'procurement/case_detail.html', context)


@login_required
def case_list(request):
    role = request.user.role
    STAGE_MAP = {
        'ACCOUNTS_CLERK': ['SURVEY', 'BID_BOQ', 'GEM_ORDER', 'CRAC', 'CRV'],
        'ACCOUNTS_JCO': ['INSPECTION'],
        'ACCOUNTS_OFFICER': ['NOTING'],
        'CFA': ['APPROVAL', 'EAS', 'SANCTION'],
    }
    if role in STAGE_MAP:
        cases = ProcurementCase.objects.filter(current_stage__in=STAGE_MAP[role])
    else:
        cases = ProcurementCase.objects.all()

    return render(request, 'procurement/case_list.html', {
        'cases': cases,
        'role': role,
        'active_nav': 'procurement',
    })


@login_required
def audit_trail(request):
    from .models import CaseStageHistory
    history = CaseStageHistory.objects.select_related('case', 'performed_by').order_by('-created_at')[:100]
    return render(request, 'procurement/audit_trail.html', {
        'history': history,
        'role': request.user.role,
        'active_nav': 'audit_trail',
    })


# ============================================================
# NOTING SHEET workflow
# ============================================================

@login_required
def procurement_dashboard(request):
    """Landing page for the "Procurement" sidebar link."""
    noting_sheets = NotingSheet.objects.select_related("requirement").order_by("-created_at")

    context = {
        "noting_sheets": noting_sheets,
        "total_count": noting_sheets.count(),
        "pending_count": noting_sheets.exclude(workflow_status="APPROVED").count(),
        "completed_count": noting_sheets.filter(workflow_status="APPROVED").count(),
        "role": request.user.role,
        "active_nav": "procurement",
    }
    return render(request, "procurement/procurement_dashboard.html", context)


@login_required
def procurement_select(request):
    """
    Step 1: dropdown of APPROVED requirements that don't already have a
    Noting Sheet. Selecting one (GET ?requirement=<pk>) shows its details
    below plus a "Create Noting Sheet" button, matching your mockup.
    """
    eligible = Requirement.objects.filter(
        workflow_status="APPROVED", noting_sheet__isnull=True
    ).order_by("-created_at")

    selected = None
    requirement_id = request.GET.get("requirement")
    if requirement_id:
        selected = get_object_or_404(Requirement, pk=requirement_id, workflow_status="APPROVED")

    return render(request, "procurement/procurement_select.html", {
        "eligible": eligible,
        "selected": selected,
        "role": request.user.role,
        "active_nav": "procurement",
    })


@login_required
def noting_sheet_create(request, requirement_pk):
    """
    CHANGED: now handles NotingSheetItemFormSet alongside the header form,
    for the mockup's multi-item Item Details table. The first item row is
    seeded from the linked Requirement's item_name/quantity/estimated_cost.
    """
    requirement = get_object_or_404(Requirement, pk=requirement_pk, workflow_status="APPROVED")

    if hasattr(requirement, "noting_sheet"):
        messages.info(request, "A Noting Sheet already exists for this requirement.")
        return redirect("noting_sheet_detail", pk=requirement.noting_sheet.pk)

    # Computed once, reused for both the initial page load (so the clerk
    # sees real figures before even touching the form) and the actual
    # save below — previously this only ran inside the save branch, so
    # the fields stayed blank until after submission.
    computed_allotted, computed_released = get_fund_balance(requirement.sub_head)

    if request.method == "POST":
        form = NotingSheetForm(request.POST)
        formset = NotingSheetItemFormSet(request.POST)
        if form.is_valid() and formset.is_valid():
            noting_sheet = form.save(commit=False)
            noting_sheet.requirement = requirement
            noting_sheet.created_by = request.user
            noting_sheet.amount_allotted = computed_allotted
            noting_sheet.amount_released = computed_released

            noting_sheet.save()

            formset.instance = noting_sheet
            formset.save()

            # Expended = this sheet's own item table total. Needs a
            # second save since it depends on the items, which can't be
            # totaled until after formset.save() has written them.
            noting_sheet.amount_expended = noting_sheet.total_amount
            noting_sheet.save(update_fields=["amount_expended"])

            messages.success(request, f"{noting_sheet.noting_id} created.")
            return redirect("noting_sheet_detail", pk=noting_sheet.pk)
    else:
        form = NotingSheetForm(initial={
            "dated": timezone.now().date(),
            "paragraph_1": f"PROCUREMENT OF {requirement.item_name.upper()} FOR {requirement.demanded_by.upper()}",
        })
        formset = NotingSheetItemFormSet(initial=[{
            "description": requirement.item_name,
            "au": "Nos",
            "quantity": requirement.quantity,
            "unit_price": requirement.estimated_cost,
        }], queryset=NotingSheetItem.objects.none())

    return render(request, "procurement/noting_sheet_form.html", {
        "form": form,
        "formset": formset,
        "requirement": requirement,
        "computed_allotted": computed_allotted,
        "computed_released": computed_released,
        "role": request.user.role,
        "active_nav": "procurement",
    })


@login_required
def noting_sheet_submit_for_approval(request, pk):
    noting_sheet = get_object_or_404(NotingSheet, pk=pk)
    if noting_sheet.created_by != request.user:
        messages.error(request, "You can only submit noting sheets you created.")
        return redirect("procurement_dashboard")

    if request.method == "POST":
        try:
            noting_sheet.submit_for_approval()
            messages.success(request, f"{noting_sheet.noting_id} sent for CFA approval.")
        except ValidationError as e:
            messages.error(request, str(e))

    return redirect("noting_sheet_detail", pk=pk)


@login_required
def noting_sheet_detail(request, pk):
    """CFA-only approval — Account Officer review is not part of this flow."""
    noting_sheet = get_object_or_404(NotingSheet, pk=pk)
    role = request.user.role

    cfa_form = None
    can_act_as_cfa = role == "CFA" and noting_sheet.workflow_status == "PENDING_CFA"

    if request.method == "POST":
        if can_act_as_cfa and "cfa_submit" in request.POST:
            cfa_form = NotingCFADecisionForm(request.POST)
            if cfa_form.is_valid():
                noting_sheet.cfa_decide(
                    user=request.user,
                    decision=cfa_form.cleaned_data["cfa_status"],
                    remarks=cfa_form.cleaned_data["cfa_remarks"],
                )
                messages.success(request, "CFA decision recorded.")
                return redirect("noting_sheet_detail", pk=pk)

    if cfa_form is None and can_act_as_cfa:
        cfa_form = NotingCFADecisionForm(initial={"cfa_status": "PENDING"})

    return render(request, "procurement/noting_sheet_detail.html", {
        "noting_sheet": noting_sheet,
        "items": noting_sheet.items.all(),
        "role": role,
        "active_nav": "procurement",
        "cfa_form": cfa_form,
        "can_submit": noting_sheet.is_editable and noting_sheet.created_by == request.user,
    })


# ============================================================
# EAS workflow — unchanged
# ============================================================

@login_required
def eas_create(request, noting_sheet_pk):
    noting_sheet = get_object_or_404(NotingSheet, pk=noting_sheet_pk, workflow_status="APPROVED")

    if hasattr(noting_sheet, "eas"):
        messages.info(request, "An EAS already exists for this noting sheet.")
        return redirect("eas_detail", pk=noting_sheet.eas.pk)

    autofill = compute_eas_autofill(noting_sheet)

    if request.method == "POST":
        form = EASForm(request.POST)
        if form.is_valid():
            eas = form.save(commit=False)
            eas.noting_sheet = noting_sheet
            eas.created_by = request.user
            for field_name, value in autofill.items():
                setattr(eas, field_name, value)
            eas.save()
            eas.submit_for_approval()
            messages.success(request, "EAS created and sent for CFA approval.")
            return redirect("eas_detail", pk=eas.pk)
    else:
        form = EASForm()

    return render(request, "procurement/eas_form.html", {
        "form": form,
        "noting_sheet": noting_sheet,
        "autofill": autofill,
        "role": request.user.role,
        "active_nav": "procurement",
    })


@login_required
def eas_edit(request, pk):
    eas = get_object_or_404(EAS, pk=pk)
    if not eas.is_editable or eas.created_by != request.user:
        messages.error(request, "This EAS can't be edited right now.")
        return redirect("eas_detail", pk=pk)

    # Recomputed live, not just read off the existing eas instance — if
    # the linked NotingSheet's data ever changed, the auto-fetched fields
    # stay in sync rather than showing a stale snapshot from creation.
    autofill = compute_eas_autofill(eas.noting_sheet)

    if request.method == "POST":
        form = EASForm(request.POST, instance=eas)
        if form.is_valid():
            eas = form.save(commit=False)
            for field_name, value in autofill.items():
                setattr(eas, field_name, value)
            eas.save()
            eas.submit_for_approval()
            messages.success(request, "EAS updated and re-sent for CFA approval.")
            return redirect("eas_detail", pk=pk)
    else:
        form = EASForm(instance=eas)

    return render(request, "procurement/eas_form.html", {
        "form": form,
        "noting_sheet": eas.noting_sheet,
        "eas": eas,
        "autofill": autofill,
        "role": request.user.role,
        "active_nav": "procurement",
    })


@login_required
def eas_detail(request, pk):
    eas = get_object_or_404(EAS, pk=pk)
    role = request.user.role

    cfa_form = None
    can_act_as_cfa = role == "CFA" and eas.workflow_status == "PENDING_CFA"

    if request.method == "POST":
        if can_act_as_cfa and "cfa_submit" in request.POST:
            cfa_form = EASCFADecisionForm(request.POST)
            if cfa_form.is_valid():
                eas.cfa_decide(
                    user=request.user,
                    decision=cfa_form.cleaned_data["cfa_status"],
                    remarks=cfa_form.cleaned_data["cfa_remarks"],
                )
                messages.success(request, "CFA decision recorded.")
                return redirect("eas_detail", pk=pk)

    if cfa_form is None and can_act_as_cfa:
        cfa_form = EASCFADecisionForm(initial={"cfa_status": "PENDING"})

    return render(request, "procurement/eas_detail.html", {
        "eas": eas,
        "noting_sheet": eas.noting_sheet,
        "role": role,
        "active_nav": "procurement",
        "cfa_form": cfa_form,
        "can_edit": eas.is_editable and eas.created_by == request.user,
    })


# ============================================================
# PDF downloads
# ============================================================

@login_required
def noting_sheet_download_pdf(request, pk):
    noting_sheet = get_object_or_404(NotingSheet, pk=pk)
    if noting_sheet.workflow_status != "APPROVED":
        messages.error(request, "This Noting Sheet can only be downloaded once it's fully approved.")
        return redirect("noting_sheet_detail", pk=pk)
    return render(request, "procurement/pdf/noting_sheet_pdf.html", {
        "noting_sheet": noting_sheet,
        "items": noting_sheet.items.all(),
    })


@login_required
def eas_download_pdf(request, pk):
    eas = get_object_or_404(EAS, pk=pk)
    if eas.workflow_status != "APPROVED":
        messages.error(request, "This EAS can only be downloaded once it's fully approved.")
        return redirect("eas_detail", pk=pk)
    return render(request, "procurement/pdf/eas_pdf.html", {"eas": eas, "noting_sheet": eas.noting_sheet})


# ============================================================
# Sanction / Contract / Invoice document uploads
# ============================================================

EAS_DOCUMENT_FIELDS = {
    "sanction": "sanction_document",
    "contract": "contract_document",
    "invoice": "invoice_document",
}


@login_required
def eas_upload_document(request, pk, doc_type):
    eas = get_object_or_404(EAS, pk=pk)
    field_name = EAS_DOCUMENT_FIELDS.get(doc_type)

    if field_name is None:
        messages.error(request, "Unknown document type.")
        return redirect("eas_detail", pk=pk)

    if eas.workflow_status != "APPROVED":
        messages.error(request, "Documents can only be uploaded once the EAS is approved.")
        return redirect("eas_detail", pk=pk)

    if request.method == "POST":
        form = EASDocumentUploadForm(request.POST, request.FILES)
        if form.is_valid():
            setattr(eas, field_name, form.cleaned_data["document"])
            eas.save()
            messages.success(request, f"{doc_type.title()} document uploaded.")
        else:
            messages.error(request, "; ".join(form.errors.get("document", ["Upload failed — PDF files only."])))

    return redirect("eas_detail", pk=pk)

# ============================================================
# CONVENING ORDER
# Simplified as per official Convening Order format
# ============================================================

CONVENING_ORDER_CREATOR_ROLES = {
    "ADMINISTRATOR",
    "HEAD_CLERK",
    "ACCOUNTS_CLERK",
}


# ============================================================
# CONVENING ORDER
# Simplified as per official Convening Order format
# ============================================================

CONVENING_ORDER_CREATOR_ROLES = {
    "ADMINISTRATOR",
    "HEAD_CLERK",
    "ACCOUNTS_CLERK",
}


@login_required
def convening_order_create(request, eas_pk):

    eas = get_object_or_404(EAS, pk=eas_pk)

    # --------------------------------------------------------
    # 1. Permission check
    # --------------------------------------------------------
    if request.user.role not in CONVENING_ORDER_CREATOR_ROLES:

        messages.error(
            request,
            "You are not authorised to create a Convening Order."
        )

        return redirect(
            "eas_detail",
            pk=eas.pk
        )

    # --------------------------------------------------------
    # 2. EAS must already be approved
    # --------------------------------------------------------
    if eas.workflow_status != "APPROVED":

        messages.error(
            request,
            "Convening Order can only be created after EAS approval."
        )

        return redirect(
            "eas_detail",
            pk=eas.pk
        )

    # --------------------------------------------------------
    # 3. Mandatory GeM documents must exist
    # --------------------------------------------------------
    if not (
        eas.sanction_document
        and eas.contract_document
        and eas.invoice_document
    ):

        messages.error(
            request,
            "Sanction, Contract and Invoice must be uploaded "
            "before creating the Convening Order."
        )

        return redirect(
            "eas_detail",
            pk=eas.pk
        )

    # --------------------------------------------------------
    # 4. Find linked Procurement Case
    # --------------------------------------------------------
    try:

        procurement_case = (
            eas
            .noting_sheet
            .requirement
            .procurement_case
        )

    except AttributeError:

        procurement_case = None

    if procurement_case is None:

        messages.error(
            request,
            "No Procurement Case is linked to this EAS."
        )

        return redirect(
            "eas_detail",
            pk=eas.pk
        )

    # --------------------------------------------------------
    # 5. Only one Convening Order per Procurement Case
    # --------------------------------------------------------
    try:

        existing_order = (
            procurement_case
            .convening_order
        )

    except ConveningOrder.DoesNotExist:

        existing_order = None

    if existing_order:

        return redirect(
            "convening_order_detail",
            pk=existing_order.pk
        )

    # --------------------------------------------------------
    # 6. Save submitted form
    # --------------------------------------------------------
    if request.method == "POST":

<<<<<<< HEAD
        form = ConveningOrderForm(
            request.POST
        )
=======
        form = ConveningOrderForm(request.POST)
>>>>>>> prince_dev

        if form.is_valid():

            order = form.save(
                commit=False
            )

<<<<<<< HEAD
=======
            # Link to Procurement Case
>>>>>>> prince_dev
            order.procurement_case = (
                procurement_case
            )

<<<<<<< HEAD
=======
            # Store the user who created it
>>>>>>> prince_dev
            order.created_by = (
                request.user
            )

<<<<<<< HEAD
=======
            # Save Convening Order
>>>>>>> prince_dev
            order.save()

            messages.success(
                request,
                f"{order.convening_order_id} "
                "created successfully."
            )

            return redirect(
                "convening_order_detail",
                pk=order.pk
            )
<<<<<<< HEAD
        else:
            print("CONVENING ORDER FORM ERRORS:")
            print(form.errors)
=======

        else:

            print(
                "CONVENING ORDER FORM ERRORS:"
            )

            print(
                form.errors
            )
>>>>>>> prince_dev

    # --------------------------------------------------------
    # 7. Initial form
    # --------------------------------------------------------
    else:

        form = ConveningOrderForm(
            initial={
                "order_date": timezone.localdate(),
                "two_ic_appointment": "2IC",
            }
        )

    # --------------------------------------------------------
    # 8. Show form
    # --------------------------------------------------------
    return render(
        request,
        "procurement/convening_order_form.html",
        {
            "form": form,

            "eas": eas,

            "procurement_case": procurement_case,

            "requirement": (
<<<<<<< HEAD
                eas.noting_sheet.requirement
=======
                eas
                .noting_sheet
                .requirement
>>>>>>> prince_dev
            ),

            # Auto fetched from EAS
            "file_no": eas.file_no,

            "role": request.user.role,

            "active_nav": "procurement",
        }
    )
<<<<<<< HEAD


@login_required
def convening_order_detail(request, pk):

    order = get_object_or_404(
        ConveningOrder.objects.select_related(
            "procurement_case",
            "procurement_case__requirement_item",
        ),
        pk=pk,
    )

    return render(
        request,
        "procurement/convening_order_detail.html",
        {
            "order": order,
            "role": request.user.role,
            "active_nav": "procurement",
        }
    )


@login_required
=======
@login_required
def convening_order_detail(request, pk):

    order = get_object_or_404(
        ConveningOrder.objects.select_related(
            "procurement_case",
            "procurement_case__requirement_item",
        ),
        pk=pk,
    )

    return render(
        request,
        "procurement/convening_order_detail.html",
        {
            "order": order,
            "role": request.user.role,
            "active_nav": "procurement",
        }
    )


@login_required
>>>>>>> prince_dev
def convening_order_download_docx(request, pk):

    order = get_object_or_404(
        ConveningOrder,
        pk=pk
    )

    # --------------------------------------------------------
    # Create Word document
    # --------------------------------------------------------

    document = Document()

    normal_style = document.styles["Normal"]

    normal_style.font.name = (
        "Times New Roman"
    )

    normal_style.font.size = Pt(12)

    # --------------------------------------------------------
    # TITLE
    # --------------------------------------------------------

    title = document.add_paragraph()

    title.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    title_run = title.add_run(
        "CONVENING ORDER"
    )

    title_run.bold = True
    title_run.underline = True

    # --------------------------------------------------------
    # PARA 1
    # --------------------------------------------------------

    para1 = document.add_paragraph()

    number_run = para1.add_run(
        "1. "
    )

    number_run.bold = True

    para1.add_run(
        order.description.strip()
    )

    # --------------------------------------------------------
    # BOARD MEMBERS
    # --------------------------------------------------------

    document.add_paragraph(
        f"Presiding Officer : "
        f"{order.presiding_officer}"
    )

    document.add_paragraph(
        f"Members             1. "
        f"{order.member_1}"
    )

    document.add_paragraph(
        f"                    2. "
        f"{order.member_2}"
    )

    # --------------------------------------------------------
    # PARA 2
    # --------------------------------------------------------

    para2 = document.add_paragraph()

    para2_number = para2.add_run(
        "2. "
    )

    para2_number.bold = True

    completion_date = (
        order
        .completion_date
        .strftime("%d %b %Y")
    )

    para2.add_run(
        "The bd proceedings duly completed "
        "in all respect will be submitted "
        f"to this HQ by {completion_date}."
    )

    # --------------------------------------------------------
    # 2IC SIGNATURE BLOCK
    # --------------------------------------------------------

    signature = document.add_paragraph()

    signature.alignment = (
        WD_ALIGN_PARAGRAPH.RIGHT
    )

    signature.add_run(
        f"({order.two_ic_name})\n"
        f"{order.two_ic_rank}\n"
        f"{order.two_ic_appointment}"
    )

    # --------------------------------------------------------
    # FILE NUMBER
    # --------------------------------------------------------

    document.add_paragraph(
        order.file_no
        or "FILE NO NOT AVAILABLE"
    )

    # --------------------------------------------------------
    # STATIC UNIT ADDRESS
    # --------------------------------------------------------

    document.add_paragraph(
        "21 SATA Regt (Plains)"
    )

    document.add_paragraph(
        "PIN : 925721"
    )

    document.add_paragraph(
        "c/o 56 APO"
    )

    # --------------------------------------------------------
    # DATE
    # --------------------------------------------------------

    order_date = (
        order
        .order_date
        .strftime("%d %b %Y")
    )

    document.add_paragraph(
        order_date
    )

    # --------------------------------------------------------
    # DISTRIBUTION
    # --------------------------------------------------------

    document.add_paragraph("")

    document.add_paragraph(
        "Distr :-"
    )

    document.add_paragraph(
        "Normal"
    )

    # --------------------------------------------------------
    # SEND WORD FILE TO USER
    # --------------------------------------------------------

    output = BytesIO()

    document.save(output)

    output.seek(0)

    filename = (
        f"{order.convening_order_id}.docx"
    )

    response = HttpResponse(
        output.getvalue(),
        content_type=(
            "application/vnd."
            "openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
    )

    response["Content-Disposition"] = (
        f'attachment; filename="{filename}"'
    )

    return response

    order = get_object_or_404(
        ConveningOrder.objects.select_related(
            "procurement_case",
            "presiding_officer",
            "member_1",
            "member_2",
            "technical_representative",
            "member_secretary",
            "convening_authority",
            "report_submission_to",
        ),
        pk=pk,
    )

    return render(request, "procurement/convening_order_detail.html", {
        "order": order,
        "role": request.user.role,
        "active_nav": "procurement",
    })